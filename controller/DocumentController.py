from docx import Document
from docx.document import Document as Doc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt

def write_card_name(document: Doc, card_name):
    p = document.add_paragraph()
    text = p.add_run(card_name)
    text.bold = True
    text.font.size = Pt(12)

def write_description(document: Doc, card_description: str):
    card_description = card_description.split('---')[-1].replace('`', '')
    document.add_paragraph().add_run('Descrição:').bold = True

    next_bold = card_description.find('**')
    while next_bold != -1:
        substr = card_description[:next_bold]
        for text in substr.split('\n'):
            clean_text = text.replace('\n', '').strip()
            if len(clean_text) < 3: continue

            p = document.add_paragraph()
            if text.startswith('   -'):
                p.add_run(text[5:])
                p.style = 'List Bullet'
            else:
                p.add_run(text).bold = True
        
        card_description = card_description[next_bold+2:]
        next_bold = card_description.find('**')

def write_info(document: Doc, title: str, info: str):
    p = document.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(info)

def write_info_list(document: Doc, title: str, info: list):
    p = document.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(str(info).replace('[', '').replace(']', '').replace("'", ''))

def write_activities(document: Doc, activities: list):
    p = document.add_paragraph()
    p.add_run('Atividades:').bold = True

    if activities:
        cont = 0
        for activity in activities:
            cont+=1
            p.add_run(f'\n\n{cont}. ').bold = True
            p.add_run(f'{activity["name"]} ({activity["state"]})')
    else:
        write_not_added(p, 'Nenhuma atividade adicionado')

def write_evidences(document: Doc, relative_path_evidences: str, evidences: list):
    images_not_found = []
    if evidences:
        for evidence in evidences:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            file_extension = str(evidence['url']).split('.')[-1].lower()
            file_name = str(evidence['name']).split('.')[0]
            file_name_with_extension = file_name + '.' + file_extension
            
            try:
                p.add_run().add_picture(f"{relative_path_evidences}/{file_name_with_extension}", width=Inches(6))
                document.add_paragraph(f'Figura: {file_name.replace("ev_","").replace("_", " ")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                text = p.add_run(f'Não foi possível baixar evidência: {file_name_with_extension}')
                text.font.color.rgb = RGBColor(255, 0, 0)
                images_not_found.append(file_name_with_extension)
    else:
        p = document.add_paragraph()
        write_not_added(p, 'Nenhuma evidência adicionada')

    if images_not_found:
        print('Imagens não encontradas: ')
        print(images_not_found)

def write_blank_line(document: Doc):
    document.add_paragraph('\n')

def write_not_added(paragraph, info):
    text = paragraph.add_run(f' {info}')
    text.italic = True
    text.font.color.rgb = RGBColor(255, 0, 0)